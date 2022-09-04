
from django.shortcuts import render,redirect
import googletrans
from django.conf import settings
import googletrans
from django.core.files.storage import FileSystemStorage
from django.core.files import File
import docx
import os
from django.contrib import messages
def index(request):
    try:
        if request.method == 'POST':
            input = request.POST.get('text_area')
            src = request.POST.get('src')
            dest = request.POST.get('dest')
            item = googletrans.LANGUAGES.values()
            
            translator = googletrans.Translator()
            out = translator.translate(input , dest=dest,src=src)

            params = {'items':item,'Input':input,'Output':out.text}
            return render(request,'text_area.html',params)

        else:

            item = googletrans.LANGUAGES.values()
        
            params = {'items':item}
            return render(request,'text_area.html',params)
    except:
        item = googletrans.LANGUAGES.values()

        params = {'items':item}
        return render(request,'text_area.html',params)

def txt_file(file,src,dest):
    
    f = file.read()
    f = f.decode('utf-8')
    new = open(os.path.join(os.getcwd(),'media')+"/"+file.name, "w")
    translator = googletrans.Translator()
    
    out = translator.translate(f, dest=dest,src=src)
    
       
    new.write(out.text)
   
    file_name = os.path.join(os.getcwd(),'media')+"/"+file.name
    
   
    return file_name
   
def word_file(file,src,dest):
    translator = googletrans.Translator()
       
    fs = FileSystemStorage()
    doc = docx.Document(file)
   
    mydoc = docx.Document()
       
    
    all_paras = doc.paragraphs

    for para in all_paras:
        a= para.text
        if a.isalpha():
        
            out = translator.translate(a, dest=dest,src=src)
            mydoc.add_paragraph(out.text)
                
            
        else:
            mydoc.add_paragraph(a)
        mydoc.save(os.path.join(os.getcwd(),'media')+"/"+file.name) 
        file_name = os.path.join(os.getcwd(),'media')+"/"+file.name
    return file_name

def file(request):
    
    if request.method == "POST" and request.FILES["file"]:
        
        file = request.FILES["file"]
        src = request.POST.get('src')
        dest = request.POST.get('dest')

       
        file_name = ""
        # Task get dest and src 

        item = googletrans.LANGUAGES.values() # Languages
        fs = FileSystemStorage()

        if file.name.split(".")[1] == 'txt':
            file_name = txt_file(file,src,dest)
        else:
            file_name = word_file(file,src,dest)
        
        if file_name == "":

            params = {'items':item,"docs": file_url}
            return render(request, "upload.html",params)

        else:
            file_url = fs.url(file_name)
            request.session['file']  = file_name
        
            messages.success(request, 'Translation Successful Click The Download Button Below.')
            params = {'items':item,"docs": file_url}
            return render(request, "upload.html",params)
    
    
    else:
        if request.session.get('file') != None:
        
            
            fs = FileSystemStorage()
            file_name = request.session.get('file')
            del request.session['file']
            os.remove(file_name)
            
            item = googletrans.LANGUAGES.values() 
            params = {'items':item}
            return render(request, "upload.html",params)
        else:
            item = googletrans.LANGUAGES.values() 
            params = {'items':item}
            return render(request, "upload.html",params)

        



    
